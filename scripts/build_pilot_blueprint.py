from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path(r"C:\Users\hp\AppData\Local\Temp\Agentic_Target_Modeling_Pilot_Implementation_Blueprint_v0.1.docx")

NAVY = "17365D"
BLUE = "2F75B5"
TEAL = "0F6B78"
LIGHT_BLUE = "D9EAF7"
LIGHT_TEAL = "DDEBF7"
LIGHT_GREY = "F2F2F2"
AMBER = "FFF2CC"
WHITE = "FFFFFF"
TEXT = RGBColor(31, 31, 31)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=110, bottom=90, end=110):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    marker = OxmlElement("w:tblHeader")
    marker.set(qn("w:val"), "true")
    tr_pr.append(marker)


def set_table_borders(table, color="B7C9DA"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = qn(f"w:{edge}")
        element = borders.find(tag)
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_width(cell, inches):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(inches * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def add_text(cell, text, bold=False, color=None, size=9):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Aptos"
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color) if color else TEXT
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    set_table_borders(table)
    header = table.rows[0]
    set_repeat_header(header)
    for i, label in enumerate(headers):
        set_width(header.cells[i], widths[i])
        set_cell_shading(header.cells[i], NAVY)
        add_text(header.cells[i], label, bold=True, color=WHITE, size=8.5)
    for row_values in rows:
        row = table.add_row()
        for i, value in enumerate(row_values):
            set_width(row.cells[i], widths[i])
            add_text(row.cells[i], str(value), size=8.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    p.add_run(text)
    return p


def add_body(doc, text, bold_lead=None):
    p = doc.add_paragraph(style="Body Text")
    if bold_lead and text.startswith(bold_lead):
        p.add_run(bold_lead).bold = True
        p.add_run(text[len(bold_lead):])
    else:
        p.add_run(text)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_note(doc, title, text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    set_width(table.rows[0].cells[0], 6.35)
    set_cell_shading(table.rows[0].cells[0], LIGHT_BLUE)
    set_table_borders(table, color="9CC2E5")
    cell = table.rows[0].cells[0]
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(title + ": ")
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(NAVY)
    p.add_run(text)
    set_cell_margins(cell, top=120, start=150, bottom=120, end=150)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.add_run("Page ")
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    paragraph._p.append(fld_char1)
    paragraph._p.append(instr_text)
    paragraph._p.append(fld_char2)


def configure_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10)
    normal.font.color.rgb = TEXT
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.08

    body = styles["Body Text"]
    body.font.name = "Aptos"
    body.font.size = Pt(10)
    body.font.color.rgb = TEXT
    body.paragraph_format.space_after = Pt(6)
    body.paragraph_format.line_spacing = 1.08

    for level, size, color in ((1, 16, NAVY), (2, 12, TEAL), (3, 10.5, NAVY)):
        style = styles[f"Heading {level}"]
        style.font.name = "Aptos Display" if level == 1 else "Aptos"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(14 if level == 1 else 10)
        style.paragraph_format.space_after = Pt(5)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = "Aptos"
        style.font.size = Pt(10)
        style.paragraph_format.space_after = Pt(3)

    header = section.header.paragraphs[0]
    header.text = "AGENTIC TARGET MODELING | PILOT IMPLEMENTATION BLUEPRINT"
    header.runs[0].font.name = "Aptos"
    header.runs[0].font.size = Pt(8)
    header.runs[0].font.color.rgb = RGBColor.from_string(TEAL)
    footer = section.footer.paragraphs[0]
    add_page_number(footer)
    for run in footer.runs:
        run.font.name = "Aptos"
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor.from_string("666666")


def main():
    doc = Document()
    configure_document(doc)

    # Opening block
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(42)
    p.paragraph_format.space_after = Pt(6)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run("PILOT IMPLEMENTATION BLUEPRINT")
    r.font.name = "Aptos Display"
    r.font.size = Pt(25)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(NAVY)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(16)
    r = p.add_run("Agentic Target Modeling Solution")
    r.font.name = "Aptos"
    r.font.size = Pt(15)
    r.font.color.rgb = RGBColor.from_string(TEAL)
    add_note(doc, "Purpose", "Translate the approved logical design into a bounded, reviewable first implementation. This blueprint assumes a read-only, recommendation-only pilot on Databricks.")
    add_table(doc, ["Document status", "Value"], [
        ("Version", "0.1 — implementation planning draft"),
        ("Primary pilot outputs", "Source observation dictionary, ontology mappings, Silver ODS recommendation, Gold recommendation, STTM package"),
        ("Pilot authority", "Recommendations only; human approval is required before an artifact is treated as authoritative"),
        ("Companion design", "Agentic Target Modeling Architecture Design v0.2"),
    ], [1.8, 4.55])

    add_heading(doc, "1. Pilot objective", 1)
    add_body(doc, "Prove that governed source evidence and approved P&C knowledge can produce a reviewable source data dictionary and a small set of trusted target-model recommendations for one bounded source/domain combination.")
    add_bullets(doc, [
        "Start with one source product/module/version and one P&C business domain or LOB.",
        "Use metadata and approved profiling evidence only; do not ingest, transform, deploy, or modify source data.",
        "Measure mapping quality, evidence completeness, reviewer effort, workflow reliability, cost, and elapsed time.",
        "Treat the pilot as a decision-quality evaluation, not as a production migration implementation.",
    ])

    add_heading(doc, "2. Decisions required before build", 1)
    add_table(doc, ["Decision", "Owner", "Pilot decision to record"], [
        ("Source boundary", "Data architect", "Source system, product/module/version, candidate objects, domain/LOB, and exclusions."),
        ("Evidence boundary", "Privacy steward + platform owner", "Allowed metadata, profiles, sample classes, sanitization rules, and retention period."),
        ("Approval model", "Architecture lead", "Named architect, domain SME/steward, privacy steward, and decision turnaround target."),
        ("Confidence policy", "Architecture lead + evaluation owner", "Score components, calibration set, review thresholds, and mandatory-review cases."),
        ("Operating model", "Platform owner", "Scheduled batch, event-driven workflow, or request/response entry point for the pilot."),
        ("Success targets", "Sponsor", "Quality, cost, latency, reviewer-throughput, and evidence-completeness thresholds."),
    ], [1.35, 1.45, 3.55])

    add_heading(doc, "3. Pilot solution architecture", 1)
    add_body(doc, "The pilot has four separable planes. The agent system can read from the governed knowledge and evidence plane, write only recommendation artifacts to the repository, and route material decisions to people.")
    add_table(doc, ["Plane", "Pilot components", "Allowed responsibility"], [
        ("Knowledge and evidence", "Unity Catalog, governed Bronze metadata, Databricks SQL profiles, approved COTS documentation, P&C ontology, enterprise standards", "Provide scoped, read-only context and evidence."),
        ("Agent reasoning", "Orchestrator; Source Intelligence; Ontology & Semantic; Silver ODS; Gold Data Product; STTM & Business Rule; Trust & Review", "Generate structured recommendations and explain evidence."),
        ("Control plane", "Lakeflow Workflows, MLflow evaluation, tool-access policy, run-state store", "Run idempotent tasks, capture versions, retry safely, and evaluate quality."),
        ("Governance and consumption", "Delta recommendation repository, review workbench, approval records", "Store evidence and decisions; present artifacts for human approval."),
    ], [1.35, 2.7, 2.3])
    add_note(doc, "Technical enforcement", "Agents receive read/query/retrieval tools only. They are not bound to schema-write, deployment, transformation-execution, credential-management, or ingestion tools.")

    add_heading(doc, "4. Source Intelligence Agent contract", 1)
    add_body(doc, "The Source Intelligence Agent is the first substantive pilot capability. Its job is to convert governed source evidence into a reviewable observation catalog and proposed source data dictionary; it does not create authoritative business definitions.")
    add_table(doc, ["Contract element", "Specification"], [
        ("Inputs", "Pilot scope; source-object inventory; Unity Catalog metadata; approved SQL profile results; permitted relationship evidence; COTS product/module/version context; enterprise naming and privacy standards."),
        ("Core actions", "Inspect schemas and keys; group related objects; recognize COTS patterns; identify probable customizations; infer candidate relationships; classify evidence as observed, inferred, or unresolved."),
        ("Outputs", "Source observation catalog; proposed data dictionary; source-system recognition assessment; custom-extension register; evidence links; confidence components; assumptions and open questions."),
        ("Hard limits", "No raw sensitive narrative indexing; no source writes; no target-schema deployment; no unstated inference presented as fact."),
        ("Escalation", "Route missing evidence, contradictory evidence, low-confidence semantics, unknown identifiers, and potential privacy issues to the assigned reviewer."),
        ("Quality checks", "Every proposed meaning has evidence; physical metadata is retained; confidence reasons are recorded; unresolved items are distinct from mapped items."),
    ], [1.4, 4.95])

    add_heading(doc, "5. Source data dictionary and observation-catalog schema", 1)
    add_body(doc, "The data dictionary is a governed recommendation artifact. It preserves the physical source facts and separately records proposed semantic interpretations, so reviewers can distinguish observation from inference.")
    add_table(doc, ["Field group", "Minimum fields", "Purpose"], [
        ("Identity and scope", "engagement_id, source_system, product, module, version, domain, LOB, object_name, attribute_name", "Establish the exact source and pilot boundary."),
        ("Physical observation", "physical_type, length/precision, nullability, key_role, ordinal_position, relationship evidence, profile summary", "Record observed source facts without semantic overclaim."),
        ("Semantic proposal", "proposed_business_name, definition, ontology_concept_id, synonym, domain/LOB, lifecycle state", "Capture proposed meaning and ontology mapping."),
        ("Trust and evidence", "observed_or_inferred, evidence_refs, confidence_score, component_scores, assumptions, contradictions, open_questions", "Make the recommendation reviewable and calibratable."),
        ("Governance", "privacy_class, privacy_rationale, owner, reviewer, approval_state, decision_rationale, artifact_version", "Support controlled approval and reuse."),
        ("Lineage", "source_query_or_catalog_ref, run_id, model_version, prompt_version, knowledge_pack_versions, timestamp", "Provide reproducibility and auditability."),
    ], [1.25, 3.6, 1.5])

    add_heading(doc, "6. Workflow, state, and recovery design", 1)
    add_table(doc, ["Stage", "State transition", "Recovery behavior"], [
        ("Create run", "pending → running", "Assign immutable run ID, scope version, context versions, and idempotency key."),
        ("Source observation", "running → completed / needs-clarification / failed", "Retry transient technical errors. Route insufficient evidence or sensitive-data exceptions to a reviewer."),
        ("Semantic and model recommendation", "completed → downstream task queued", "Run only when prerequisite artifact versions are valid and within pilot scope."),
        ("Trust evaluation", "completed → review-ready / remediation-required", "Return only affected artifacts to the responsible agent; record contradiction and remediation reason."),
        ("Human review", "review-ready → approved / modified / rejected / deferred", "A decision invalidates only dependent downstream artifacts. Approved unrelated artifacts remain reusable."),
        ("Reuse", "approved → governed context", "Publish approved version to the repository with scope filters; never overwrite prior decisions."),
    ], [1.35, 2.15, 2.85])
    add_note(doc, "Idempotency", "A repeated request with the same pilot scope, evidence snapshot, knowledge versions, and task contract returns or references the same run artifact rather than creating an uncontrolled duplicate.")

    add_heading(doc, "7. Confidence and review-routing policy", 1)
    add_body(doc, "Confidence is an evidence score, not an LLM self-rating. The pilot must record both the aggregate score and the component explanations used to calculate it.")
    add_table(doc, ["Score component", "Example evidence", "Routing effect"], [
        ("Retrieval fitness", "Authoritative source/version match; relevance of retrieved ontology and COTS references.", "Low fit blocks semantic progression."),
        ("Metadata-pattern strength", "Data type, key role, naming pattern, relationship and profile consistency.", "Weak pattern requires additional evidence or review."),
        ("COTS-model match", "Structure aligns to an authorized product/module/version reference model.", "Mismatch prompts customization or unresolved classification."),
        ("Rule consistency", "Proposal does not contradict enterprise standards, approved rules, or prior decisions.", "Conflict requires trust-agent remediation."),
        ("Reviewer calibration", "Comparison with labelled reviewer outcomes and observed override/rejection rates.", "Drift triggers threshold review and recalibration."),
    ], [1.55, 3.25, 1.55])
    add_bullets(doc, [
        "Mandatory human review: business-critical identifiers, financial measures, privacy classifications, candidate ontology extensions, and any contradictory or low-confidence result.",
        "Automatic progression is permitted only between analysis tasks. It never authorizes implementation, deployment, or publication as an approved enterprise model.",
    ])

    add_heading(doc, "8. Sensitive-data and tenancy controls", 1)
    add_table(doc, ["Control", "Pilot requirement"], [
        ("Evidence minimization", "Use metadata and approved aggregate profiles first. Use value-level evidence only where the evidence class is explicitly approved."),
        ("Retrieval safety", "Do not embed raw, masked, medical, legal, or narrative PII by default. Vector indexes contain approved documentation and sanitized metadata only."),
        ("Engagement isolation", "Use separate Unity Catalog permissions, scoped service identities, recommendation schemas, retrieval collections, and audit records per engagement."),
        ("Retention", "Record evidence-retention period and purge/archival procedure before collecting any value-level evidence."),
        ("Audit", "Retain run ID, data-access references, knowledge versions, reviewer decisions, and tool invocations for every material recommendation."),
    ], [1.55, 4.8])

    add_heading(doc, "9. Pilot delivery plan", 1)
    add_table(doc, ["Work package", "Primary deliverables", "Exit condition"], [
        ("1. Inception and controls", "Signed pilot scope, evidence classification, named reviewers, success targets, tool-access policy.", "Pilot boundary and controls approved."),
        ("2. Knowledge and repository foundation", "Versioned ontology/COTS/standards packs, recommendation repository schema, review-state model.", "A scoped context package can be assembled reproducibly."),
        ("3. Source Intelligence MVP", "Read-only metadata connector/query set, observation catalog, proposed data dictionary, reviewer queue.", "Reviewers can accept, modify, reject, or defer source-dictionary entries."),
        ("4. Modeling and mapping MVP", "Ontology mapping, Silver ODS, Gold model, STTM, and trust-evaluation contracts for selected objects.", "End-to-end recommendation package is reviewable."),
        ("5. Evaluation and decision", "Labelled test set, metrics dashboard, reviewer outcomes, cost/latency findings, go/no-go recommendation.", "Sponsor decision recorded with evidence."),
    ], [1.6, 3.1, 1.65])

    add_heading(doc, "10. First-release acceptance targets", 1)
    add_table(doc, ["Measure", "Target"], [
        ("Evidence completeness", "100% of material recommendations include evidence, provenance, assumptions, contradictions, confidence reasons, and approval state."),
        ("Sensitive-evidence compliance", "100% of retrieved/indexed inputs pass the approved evidence-classification and sanitization controls."),
        ("Mapping quality", "Threshold to be set against a labelled reviewer set before build; report precision, recall, and reviewer override rate."),
        ("Review throughput", "Target median decision time set during inception; report bottlenecks by reviewer role and artifact type."),
        ("Workflow reliability", "All task failures are traceable; retry and invalidation behavior is demonstrated using controlled test scenarios."),
        ("Cost and latency", "Project-defined budget and elapsed-time target per source object, measured by workflow stage."),
    ], [1.7, 4.65])

    add_heading(doc, "11. Immediate next action", 1)
    add_note(doc, "Recommended starting point", "Run the inception workshop and populate the pilot decision record. The first build artifact should be the source-object inventory and evidence-classification matrix for the chosen product/module/version and domain.")

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
