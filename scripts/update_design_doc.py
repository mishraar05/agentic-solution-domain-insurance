from pathlib import Path
from copy import deepcopy

from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement


SOURCE = Path(r"C:\Users\hp\Downloads\Agentic_Target_Modeling_Architecture_Design.docx")
OUTPUT = Path(r"C:\Users\hp\OneDrive\Documents\New project\Agentic_Target_Modeling_Architecture_Design_v0.2.docx")


def add_row(table, values):
    cells = table.add_row().cells
    for cell, value in zip(cells, values):
        cell.text = value


def insert_after(paragraph, new_element):
    paragraph._p.addnext(new_element)


def move_after(paragraph, item):
    element = item._element
    element.getparent().remove(element)
    insert_after(paragraph, element)
    return item


def main():
    doc = Document(SOURCE)

    # Revision metadata: preserve the original wording and make the review update explicit.
    for paragraph in doc.paragraphs:
        if paragraph.text == "Status: Draft for architecture review":
            paragraph.text = "Status: Draft for architecture review — updated with design-review controls"
        elif paragraph.text == "Version: 0.1":
            paragraph.text = "Version: 0.2"
        elif paragraph.text == "Date: 11 July 2026":
            paragraph.text = "Date: 11 July 2026 (design-review update)"

    # Expand the stated design decisions (Table 1) with enforceable controls.
    decisions = doc.tables[0]
    additions = [
        ("D-08", "Tool-access enforcement", "Agents are bound only to approved read/query/retrieval tools. No agent receives schema-write, deployment, transformation-execution, credential-management, or source-ingestion tool access."),
        ("D-09", "Confidence and progression", "Confidence is a calibrated evidence score, not an LLM self-assessment. It controls workflow routing only; no material recommendation becomes authoritative without human approval."),
        ("D-10", "Workflow reliability", "Use durable, idempotent workflow runs with artifact versioning, targeted retries, and downstream invalidation. Re-run only artifacts affected by new evidence or reviewer decisions."),
        ("D-11", "Sensitive evidence", "Do not index raw, masked, or narrative PII by default. Retrieve only approved documentation and sanitized metadata; governed profiling queries remain the preferred path for sensitive evidence."),
        ("D-12", "Engagement isolation", "Separate client engagements through Unity Catalog boundaries, scoped identities, engagement-specific retrieval collections, and isolated audit records."),
    ]
    for row in additions:
        add_row(decisions, row)

    # Expand the unresolved decisions list (Table 14) where a project choice is still required.
    remaining = doc.tables[13]
    for row in [
        ("R-08", "Which workflow topology will be used for the first implementation: event-driven durable orchestration, scheduled batch, or synchronous request/response?"),
        ("R-09", "What evidence classes may be retrieved or indexed, and what sanitization and approval controls apply to protected or narrative data?"),
        ("R-10", "What measurable quality, reviewer-throughput, cost, and latency thresholds define first-release success?"),
    ]:
        add_row(remaining, row)

    # Add a cohesive review-update section before the glossary, then move it into position.
    anchor = next(p for p in doc.paragraphs if p.text == "Appendix A — Glossary")
    heading = doc.add_paragraph("19. Design-review updates", style="Heading 1")
    intro = doc.add_paragraph(
        "This section converts the design-review observations into explicit controls. These controls refine the v0.1 logical design without changing its recommendation-only boundary."
    )

    confidence_heading = doc.add_paragraph("19.1 Confidence, review routing, and calibration", style="Heading 2")
    for text in [
        "Confidence is calculated from evidence quality and agreement: retrieval fit, metadata-pattern strength, COTS-model match, business-rule consistency, contradiction checks, and historical reviewer outcomes. An LLM may explain the score but must not self-certify it.",
        "Confidence is stored at recommendation and attribute level with its component reasons, evidence references, assumptions, and contradictions. The method is evaluated against a labelled reviewer-decision set and recalibrated when observed error rates drift.",
        "Confidence may permit automatic progression to the next analysis task when evidence is sufficient. It never permits autonomous implementation, publication as authoritative, or bypass of required architect, steward, or privacy review.",
    ]:
        doc.add_paragraph(text, style="List Bullet")

    workflow_heading = doc.add_paragraph("19.2 Workflow execution, iteration, and recovery", style="Heading 2")
    for text in [
        "The orchestrator manages durable workflow runs with a unique run ID, an immutable scope version, and idempotent task execution. Each agent writes versioned outputs and a clear task state: pending, running, completed, needs-clarification, failed, or superseded.",
        "A reviewer rejection or new evidence invalidates only the affected artifact and its dependent downstream artifacts. For example, a rejected Silver entity re-runs the relevant semantic, Silver, Gold, STTM, and trust tasks; unrelated domains and approved artifacts remain unchanged.",
        "Transient failures use bounded retries with recorded error context. Repeated failures, contradictory evidence, missing evidence, and business-critical uncertainty are routed to a named reviewer rather than silently retried or inferred.",
    ]:
        doc.add_paragraph(text, style="List Bullet")

    enforcement_heading = doc.add_paragraph("19.3 Enforcement, sensitive evidence, and engagement isolation", style="Heading 2")
    for text in [
        "Recommendation-only behavior is enforced technically through agent tool binding and least-privilege service identities. Write, deployment, execution, credential, and ingestion capabilities are absent from the agent tool surface; this restriction is auditable.",
        "Before any sensitive evidence is used for semantic inference, the project must classify it and apply approved minimization, masking, and retention controls. Raw narrative claims, medical, legal, and other protected text are not embedded or indexed by default.",
        "Concurrent engagements are isolated through Unity Catalog permissions, dedicated source and recommendation schemas, scoped retrieval collections, service identities, and audit records. Retrieval filters must include engagement, product version, domain/LOB, artifact type, and approval state.",
    ]:
        doc.add_paragraph(text, style="List Bullet")

    metrics_heading = doc.add_paragraph("19.4 Measurable first-release acceptance targets", style="Heading 2")
    metrics_intro = doc.add_paragraph(
        "The implementation team will set baseline values during inception and report the following measures by source, domain, product/module/version, and release."
    )
    metrics = doc.add_table(rows=1, cols=3)
    metrics.style = "Table Grid"
    for cell, label in zip(metrics.rows[0].cells, ["Measure", "Definition", "Release target"]):
        cell.text = label
    for row in [
        ("Evidence completeness", "Percentage of material recommendations with evidence, provenance, assumptions, contradictions, and approval status.", "100%"),
        ("Reviewer override rate", "Percentage of reviewed recommendations materially changed or rejected by an authorized reviewer.", "Baseline and reduce by release"),
        ("Mapping quality", "Precision and recall of approved source-to-ontology and source-to-target mappings against a labelled review set.", "Project-defined threshold"),
        ("Review throughput", "Median elapsed time from recommendation creation to final reviewer decision.", "Project-defined service level"),
        ("Operational efficiency", "Cost and elapsed time per source object, with failed/retried task rate.", "Project-defined budget and service level"),
        ("Sensitive-evidence compliance", "Percentage of retrieval/index inputs passing the approved evidence-classification and sanitization controls.", "100%"),
    ]:
        add_row(metrics, row)

    items = [heading, intro, confidence_heading]
    # Include all newly-created elements after the heading through the metrics table.
    start = heading._element
    end = metrics._element
    body = list(doc._body._element)
    start_index = body.index(start)
    end_index = body.index(end)
    block = body[start_index : end_index + 1]
    for element in block:
        element.getparent().remove(element)
    insert_position = anchor._p
    for element in block:
        insert_position.addprevious(element)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
